"""Per-filetype text extraction for the ingestion pipeline."""

from __future__ import annotations

from pathlib import Path

TEXT_EXTENSIONS = {".txt", ".md"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}

SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS

_MIN_EXTRACTABLE_CHARS = 20


class ExtractionError(Exception):
    """Raised when a file can't be extracted and should be routed to _quarantine."""


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        return _extract_plain_text(path)
    if suffix in PDF_EXTENSIONS:
        return _extract_pdf(path)
    if suffix in DOCX_EXTENSIONS:
        return _extract_docx(path)
    raise ExtractionError(f"Unsupported file extension: {suffix!r}")


def _extract_plain_text(path: Path) -> str:
    # utf-8-sig strips a leading BOM if present (common in files saved by Notepad,
    # PowerShell's Out-File, etc.) while behaving identically to utf-8 otherwise.
    text = path.read_text(encoding="utf-8-sig", errors="replace").strip()
    if not text:
        raise ExtractionError("File is empty")
    return text


def _extract_pdf(path: Path) -> str:
    import fitz  # PyMuPDF

    parts = []
    with fitz.open(path) as doc:
        for page in doc:
            parts.append(page.get_text())
    text = "\n".join(parts).strip()
    if len(text) < _MIN_EXTRACTABLE_CHARS:
        raise ExtractionError(
            "PDF has little or no extractable text (likely scanned; OCR is not supported in v1)"
        )
    return text


def _extract_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    text = "\n".join(parts).strip()
    if not text:
        raise ExtractionError("No extractable text found in .docx")
    return text
